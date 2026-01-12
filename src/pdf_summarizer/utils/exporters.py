"""
摘要导出模块

支持导出格式:
- JSON: 结构化数据
- Markdown: 可读性文档
- PDF: 摘要报告
- DOCX: Word 文档
"""

import json
import io
from typing import Dict, Any, Optional, Union
from datetime import datetime
from dataclasses import dataclass
from abc import ABC, abstractmethod
from loguru import logger

# 尝试导入可选依赖
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.debug("reportlab 未安装，PDF 导出功能不可用")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.debug("python-docx 未安装，DOCX 导出功能不可用")


@dataclass
class SummaryExportData:
    """导出数据结构"""
    summary: str
    word_count: int
    language: str
    original_length: int
    processing_time: float
    chunks_processed: int
    style: str = "detailed"
    strategy: str = "auto"
    is_medical: bool = False
    source_filename: Optional[str] = None
    created_at: datetime = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now()
        if self.metadata is None:
            self.metadata = {}
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "summary": self.summary,
            "word_count": self.word_count,
            "language": self.language,
            "original_length": self.original_length,
            "processing_time": self.processing_time,
            "chunks_processed": self.chunks_processed,
            "style": self.style,
            "strategy": self.strategy,
            "is_medical": self.is_medical,
            "source_filename": self.source_filename,
            "created_at": self.created_at.isoformat() if self.created_at else None,
            "metadata": self.metadata,
        }


class BaseExporter(ABC):
    """导出器基类"""
    
    @abstractmethod
    def export(self, data: SummaryExportData) -> bytes:
        """导出数据"""
        pass
    
    @property
    @abstractmethod
    def content_type(self) -> str:
        """MIME 类型"""
        pass
    
    @property
    @abstractmethod
    def file_extension(self) -> str:
        """文件扩展名"""
        pass


class JSONExporter(BaseExporter):
    """JSON 导出器"""
    
    def export(self, data: SummaryExportData) -> bytes:
        output = {
            "version": "1.0",
            "export_time": datetime.now().isoformat(),
            "data": data.to_dict(),
        }
        return json.dumps(output, ensure_ascii=False, indent=2).encode("utf-8")
    
    @property
    def content_type(self) -> str:
        return "application/json"
    
    @property
    def file_extension(self) -> str:
        return ".json"


class MarkdownExporter(BaseExporter):
    """Markdown 导出器"""
    
    def export(self, data: SummaryExportData) -> bytes:
        lines = []
        
        # 标题
        title = "医学文献摘要" if data.is_medical else "文档摘要"
        lines.append(f"# {title}")
        lines.append("")
        
        # 元信息
        lines.append("## 基本信息")
        lines.append("")
        if data.source_filename:
            lines.append(f"- **源文件**: {data.source_filename}")
        lines.append(f"- **语言**: {'中文' if data.language == 'zh' else 'English'}")
        lines.append(f"- **摘要风格**: {data.style}")
        lines.append(f"- **原文长度**: {data.original_length:,} 字符")
        lines.append(f"- **摘要字数**: {data.word_count:,}")
        lines.append(f"- **处理耗时**: {data.processing_time}s")
        lines.append(f"- **分块数**: {data.chunks_processed}")
        lines.append(f"- **处理策略**: {data.strategy}")
        if data.is_medical:
            lines.append(f"- **模式**: 医学文献专用")
        lines.append(f"- **生成时间**: {data.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("")
        
        # 摘要内容
        lines.append("## 摘要内容")
        lines.append("")
        lines.append(data.summary)
        lines.append("")
        
        # 注释
        lines.append("---")
        lines.append("")
        lines.append("*本摘要由 PDF Summarization Service 自动生成*")
        
        return "\n".join(lines).encode("utf-8")
    
    @property
    def content_type(self) -> str:
        return "text/markdown"
    
    @property
    def file_extension(self) -> str:
        return ".md"


class PDFExporter(BaseExporter):
    """PDF 导出器"""
    
    def __init__(self):
        if not PDF_AVAILABLE:
            raise ImportError("PDF 导出需要安装 reportlab: pip install reportlab")
    
    def export(self, data: SummaryExportData) -> bytes:
        buffer = io.BytesIO()
        
        # 创建 PDF 文档
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
        )
        
        # 获取样式
        styles = getSampleStyleSheet()
        
        # 自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=1,  # 居中
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=8,
        )
        
        # 构建内容
        story = []
        
        # 标题
        title = "医学文献摘要报告" if data.is_medical else "文档摘要报告"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # 基本信息表格
        story.append(Paragraph("基本信息", heading_style))
        
        info_data = [
            ["源文件", data.source_filename or "未知"],
            ["语言", "中文" if data.language == "zh" else "English"],
            ["摘要风格", data.style],
            ["原文长度", f"{data.original_length:,} 字符"],
            ["摘要字数", f"{data.word_count:,}"],
            ["处理耗时", f"{data.processing_time}s"],
            ["生成时间", data.created_at.strftime("%Y-%m-%d %H:%M:%S")],
        ]
        
        table = Table(info_data, colWidths=[4*cm, 10*cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(table)
        story.append(Spacer(1, 20))
        
        # 摘要内容
        story.append(Paragraph("摘要内容", heading_style))
        
        # 分段处理摘要
        paragraphs = data.summary.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # 处理特殊字符
                safe_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_para, body_style))
        
        # 页脚说明
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=1,
        )
        story.append(Paragraph(
            "本摘要由 PDF Summarization Service 自动生成",
            footer_style
        ))
        
        # 生成 PDF
        doc.build(story)
        
        buffer.seek(0)
        return buffer.read()
    
    @property
    def content_type(self) -> str:
        return "application/pdf"
    
    @property
    def file_extension(self) -> str:
        return ".pdf"


class DOCXExporter(BaseExporter):
    """DOCX 导出器"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX 导出需要安装 python-docx: pip install python-docx")
    
    def export(self, data: SummaryExportData) -> bytes:
        document = Document()
        
        # 标题
        title = "医学文献摘要报告" if data.is_medical else "文档摘要报告"
        heading = document.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        document.add_heading("基本信息", level=1)
        
        info_items = [
            ("源文件", data.source_filename or "未知"),
            ("语言", "中文" if data.language == "zh" else "English"),
            ("摘要风格", data.style),
            ("原文长度", f"{data.original_length:,} 字符"),
            ("摘要字数", f"{data.word_count:,}"),
            ("处理耗时", f"{data.processing_time}s"),
            ("分块数", str(data.chunks_processed)),
            ("处理策略", data.strategy),
            ("生成时间", data.created_at.strftime("%Y-%m-%d %H:%M:%S")),
        ]
        
        # 添加表格
        table = document.add_table(rows=len(info_items), cols=2)
        table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(info_items):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
            # 设置第一列加粗
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        document.add_paragraph()
        
        # 摘要内容
        document.add_heading("摘要内容", level=1)
        
        # 分段添加摘要
        paragraphs = data.summary.split('\n\n')
        for para in paragraphs:
            if para.strip():
                document.add_paragraph(para.strip())
        
        # 页脚
        document.add_paragraph()
        footer = document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("本摘要由 PDF Summarization Service 自动生成")
        run.italic = True
        run.font.size = Pt(9)
        
        # 保存到内存
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        return buffer.read()
    
    @property
    def content_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    @property
    def file_extension(self) -> str:
        return ".docx"


# ===========================================
# 导出器工厂
# ===========================================

EXPORTERS = {
    "json": JSONExporter,
    "markdown": MarkdownExporter,
    "md": MarkdownExporter,
}

# 可选导出器
if PDF_AVAILABLE:
    EXPORTERS["pdf"] = PDFExporter

if DOCX_AVAILABLE:
    EXPORTERS["docx"] = DOCXExporter
    EXPORTERS["word"] = DOCXExporter


def get_exporter(format: str) -> BaseExporter:
    """
    获取导出器实例
    
    Args:
        format: 导出格式 (json/markdown/pdf/docx)
        
    Returns:
        导出器实例
        
    Raises:
        ValueError: 不支持的格式
    """
    format_lower = format.lower()
    
    if format_lower not in EXPORTERS:
        available = list(EXPORTERS.keys())
        raise ValueError(f"不支持的导出格式: {format}，可用格式: {available}")
    
    return EXPORTERS[format_lower]()


def export_summary(
    data: Union[SummaryExportData, Dict[str, Any]],
    format: str = "json",
) -> tuple[bytes, str, str]:
    """
    导出摘要
    
    Args:
        data: 摘要数据
        format: 导出格式
        
    Returns:
        (文件内容, MIME类型, 文件扩展名)
    """
    # 转换字典到数据类
    if isinstance(data, dict):
        data = SummaryExportData(**data)
    
    exporter = get_exporter(format)
    content = exporter.export(data)
    
    return content, exporter.content_type, exporter.file_extension


def list_available_formats() -> list[str]:
    """列出可用的导出格式"""
    return list(EXPORTERS.keys())
